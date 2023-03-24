from django.shortcuts import redirect, render
from django.http.response import HttpResponse
from django.core.files.storage import FileSystemStorage
from django.contrib import messages

from .utils import *
from .models import summarizer, asr, action_point_classifier

from pathlib import Path
import itertools
import os
import docx
import mimetypes

def summary(request):
    
    if request.method == 'POST':
        file = request.FILES['meeting_file']

        max_slp = request.POST.get('max_sum_len', 70) # max_summary_len_percent
        if max_slp == "":
            max_slp = 0.7
        else:
            max_slp = int(max_slp) / 100

        min_slp = request.POST.get('min_sum_len', 30)
        if min_slp == "":
            min_slp = 0.3
        else:    
            min_slp = int(min_slp) / 100 # min_summary_len_percent
        
        doc_extensions = ['doc', 'docx']
        audio_extensions = ['mp3', 'wav']
        video_extensions = ['mp4']
        suported_extensions = list(itertools.chain(
                                                    doc_extensions,
                                                    audio_extensions,
                                                    video_extensions,
                                                   ))
        extension = file.name.split('.')[-1]
        
        if extension in suported_extensions:
            
            save_dir = f'summarizer/data/{extension}/'
            fs = FileSystemStorage(location=save_dir)
            f = fs.save(file.name, file)
            file_location = save_dir + f

            if extension in doc_extensions:
                meet = TeamsMeet.from_doc(file_location)
                meeting_len = meet.duration()
                num_speakers = str(meet.num_speakers())
                action_points = action_point_classifier.get_action_points(meet)
                for ap in action_points:
                    messages.add_message(request, messages.INFO, ap)
                summary = summarizer.summarize(meet.as_str(), max_slp, min_slp)

            elif extension in audio_extensions:

                if extension == 'mp3':
                    Path('summarizer/data/wav/').mkdir(exist_ok=True, parents=True)
                    wav_file_path = 'summarizer/data/wav/' + os.path.splitext(f)[0] + '.wav'
                    mp3_to_wav(file_location, wav_file_path)
                
                elif extension == 'wav':
                    wav_file_path = file_location
                

                text = asr.recognize(wav_file_path)[0]
                meeting_len = get_meeting_length_from_audio(wav_file_path)
                num_speakers = "Number of speakers can only be found using transcripts."
                action_points = action_point_classifier.get_action_points(text)
                for ap in action_points:
                    messages.add_message(request, messages.INFO, ap)
                summary = summarizer.summarize(text_splitter(text), max_slp, min_slp)

            elif extension in video_extensions:
                Path('summarizer/data/wav/').mkdir(exist_ok=True, parents=True)
                wav_file_path = 'summarizer/data/wav/' + os.path.splitext(f)[0] + '.wav'
                video_to_audio(file_location, wav_file_path)
                text = asr.recognize(wav_file_path)[0]
                meeting_len = get_meeting_length_from_audio(wav_file_path)
                num_speakers = "Number of speakers can only be found for transcripts."
                action_points = action_point_classifier.get_action_points(text)
                for ap in action_points:
                    messages.add_message(request, messages.INFO, ap)
                summary = summarizer.summarize(text_splitter(text) , max_slp, min_slp)

            Path('summarizer/data/generated/').mkdir(exist_ok=True, parents=True)
            doc =docx.Document()
            doc.add_heading('STM 4', 0)
            doc.add_paragraph("Meeting Duration: " + meeting_len)
            doc.add_paragraph("Number of Participants: " + num_speakers)
            doc.add_paragraph("Action Points:")
            for ap in action_points:
                doc.add_paragraph(ap, style='List Bullet 2')
            doc.add_paragraph("Summary:")
            doc.add_paragraph(summary)
            summary_doc_path = 'summarizer/data/generated/' + f
            doc.save(summary_doc_path)

            request.session['summary_doc_path'] = summary_doc_path

            return render(request, 'summary.html', {'summary': summary, 'meeting_len': meeting_len, 'num_speakers': num_speakers, 'action_points': action_points})

        else:
            data = f'File Extension not supported please uploat from {suported_extensions}'
            return render(request, 'index.html', {'summary': data})
    
    else: 
        return render(request, 'index.html')

def download(request):

    summary_doc_path = request.session['summary_doc_path']
    path = open(summary_doc_path, 'rb')
    mime_type, _ = mimetypes.guess_type(summary_doc_path)
    response = HttpResponse(path, content_type=mime_type)
    print("File path", summary_doc_path)
    response['Content-Disposition'] = f"attachment; filename=summary.docx"
    return response
    
